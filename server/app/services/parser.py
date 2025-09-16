import fitz  # PyMuPDF
from docx import Document
import os
from typing import Optional

class ResumeParser:
    @staticmethod
    def extract_text(file_path: str) -> Optional[str]:
        """Extract text from PDF, DOCX, or TXT files"""
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension == '.pdf':
                return ResumeParser._extract_from_pdf(file_path)
            elif file_extension == '.docx':
                return ResumeParser._extract_from_docx(file_path)
            elif file_extension == '.txt':
                return ResumeParser._extract_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            print(f"Error extracting text from {file_path}: {str(e)}")
            return None
    
    @staticmethod
    def _extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF using PyMuPDF"""
        text = ""
        with fitz.open(file_path) as doc:
            for page in doc:
                text += page.get_text()
        return text.strip()
    
    @staticmethod
    def _extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX using python-docx"""
        doc = Document(file_path)
        text = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        
        return "\n".join(text)
    
    @staticmethod
    def _extract_from_txt(file_path: str) -> str:
        """Extract text from TXT file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read().strip()